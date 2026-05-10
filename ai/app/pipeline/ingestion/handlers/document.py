"""
Atopsy — Document File Handler.

Handles: PDF, DOCX, TXT, CSV, XLSX
Extracts: author, creation date, page count, encoding, content preview.
"""

from __future__ import annotations

import os
from typing import Any

from app.core.logger import logger
from app.exceptions.pipeline import (
    FileValidationError,
    MetadataExtractionError,
)
from app.pipeline.ingestion.handlers.base import FileHandler


class DocumentHandler(FileHandler):
    """Handles document file types: PDF, DOCX, TXT, CSV, XLSX."""

    @property
    def supported_mime_types(self) -> set[str]:
        return {
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
            "text/csv",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        }

    def validate(
        self, file_path: str, mime_type: str
    ) -> list[str]:
        """Validate document integrity."""
        warnings: list[str] = []

        if not os.path.exists(file_path):
            raise FileValidationError(
                "File does not exist on disk",
                context={"file_path": file_path},
            )

        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise FileValidationError(
                "File is empty (0 bytes)",
                context={"file_path": file_path},
            )

        # PDF-specific validation
        if mime_type == "application/pdf":
            try:
                with open(file_path, "rb") as f:
                    header = f.read(5)
                if header != b"%PDF-":
                    warnings.append(
                        "File does not start with PDF magic bytes"
                    )
            except Exception as e:
                warnings.append(f"PDF header check failed: {e}")

        # DOCX validation (ZIP-based)
        if "wordprocessingml" in mime_type:
            try:
                import zipfile

                if not zipfile.is_zipfile(file_path):
                    warnings.append("DOCX file is not a valid ZIP archive")
            except Exception as e:
                warnings.append(f"DOCX validation failed: {e}")

        return warnings

    def extract_metadata(
        self, file_path: str, mime_type: str
    ) -> dict[str, Any]:
        """Route to the appropriate metadata extractor."""
        try:
            if mime_type == "application/pdf":
                return self._extract_pdf_metadata(file_path)
            elif "wordprocessingml" in mime_type:
                return self._extract_docx_metadata(file_path)
            elif mime_type == "text/plain":
                return self._extract_text_metadata(file_path)
            elif mime_type == "text/csv":
                return self._extract_csv_metadata(file_path)
            elif "spreadsheet" in mime_type or "excel" in mime_type:
                return self._extract_xlsx_metadata(file_path)
            else:
                return {"meta_type": "document", "raw": {}}

        except MetadataExtractionError:
            raise
        except Exception as e:
            logger.warning(
                f"Metadata extraction failed for {file_path}: {e}"
            )
            return {
                "meta_type": "document",
                "error": str(e),
                "raw": {},
            }

    # ── PDF ──────────────────────────────────

    def _extract_pdf_metadata(
        self, file_path: str
    ) -> dict[str, Any]:
        metadata: dict[str, Any] = {"meta_type": "document"}

        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)
                info = pdf.metadata or {}
                metadata["author"] = info.get("Author")
                metadata["creator"] = info.get("Creator")
                metadata["producer"] = info.get("Producer")
                metadata["title"] = info.get("Title")
                metadata["subject"] = info.get("Subject")
                metadata["creation_date"] = info.get("CreationDate")
                metadata["modification_date"] = info.get("ModDate")

                # Extract first page text as preview
                if pdf.pages:
                    first_text = pdf.pages[0].extract_text()
                    if first_text:
                        metadata["content_preview"] = first_text[:1000]

        except ImportError:
            logger.warning("pdfplumber not installed, trying PyMuPDF")
            try:
                import fitz  # PyMuPDF

                doc = fitz.open(file_path)
                metadata["page_count"] = doc.page_count
                info = doc.metadata or {}
                metadata["author"] = info.get("author")
                metadata["title"] = info.get("title")
                metadata["subject"] = info.get("subject")
                metadata["creator"] = info.get("creator")
                metadata["creation_date"] = info.get("creationDate")
                metadata["modification_date"] = info.get("modDate")

                if doc.page_count > 0:
                    first_text = doc[0].get_text()
                    if first_text:
                        metadata["content_preview"] = first_text[:1000]
                doc.close()

            except ImportError:
                metadata["warning"] = (
                    "Neither pdfplumber nor PyMuPDF available"
                )

        return metadata

    # ── DOCX ─────────────────────────────────

    def _extract_docx_metadata(
        self, file_path: str
    ) -> dict[str, Any]:
        metadata: dict[str, Any] = {"meta_type": "document"}

        try:
            from docx import Document

            doc = Document(file_path)
            props = doc.core_properties

            metadata["author"] = props.author
            metadata["title"] = props.title
            metadata["subject"] = props.subject
            metadata["creation_date"] = (
                props.created.isoformat() if props.created else None
            )
            metadata["modification_date"] = (
                props.modified.isoformat() if props.modified else None
            )
            metadata["last_modified_by"] = props.last_modified_by
            metadata["revision"] = props.revision
            metadata["page_count"] = None  # Not reliable in python-docx
            metadata["paragraph_count"] = len(doc.paragraphs)

            # Content preview
            text_parts = [p.text for p in doc.paragraphs[:20] if p.text]
            metadata["content_preview"] = "\n".join(text_parts)[:1000]

        except ImportError:
            metadata["warning"] = "python-docx not installed"
        except Exception as e:
            metadata["error"] = str(e)

        return metadata

    # ── TXT ──────────────────────────────────

    def _extract_text_metadata(
        self, file_path: str
    ) -> dict[str, Any]:
        metadata: dict[str, Any] = {"meta_type": "document"}

        try:
            # Detect encoding
            with open(file_path, "rb") as f:
                raw = f.read(8192)

            import chardet

            detected = chardet.detect(raw)
            encoding = detected.get("encoding", "utf-8")
            confidence = detected.get("confidence", 0.0)

            metadata["encoding"] = encoding
            metadata["encoding_confidence"] = confidence

            with open(file_path, "r", encoding=encoding, errors="replace") as f:
                content = f.read()

            metadata["line_count"] = content.count("\n") + 1
            metadata["character_count"] = len(content)
            metadata["word_count"] = len(content.split())
            metadata["content_preview"] = content[:1000]

        except ImportError:
            # chardet not available — fall back to utf-8
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            metadata["encoding"] = "utf-8"
            metadata["line_count"] = content.count("\n") + 1
            metadata["character_count"] = len(content)
            metadata["content_preview"] = content[:1000]

        except Exception as e:
            metadata["error"] = str(e)

        return metadata

    # ── CSV ──────────────────────────────────

    def _extract_csv_metadata(
        self, file_path: str
    ) -> dict[str, Any]:
        metadata: dict[str, Any] = {"meta_type": "document"}

        try:
            import pandas as pd

            df = pd.read_csv(file_path, nrows=100)
            metadata["row_count_sample"] = len(df)
            metadata["column_count"] = len(df.columns)
            metadata["columns"] = list(df.columns)
            metadata["dtypes"] = {
                col: str(dtype) for col, dtype in df.dtypes.items()
            }

            # Full row count (without loading all data)
            with open(file_path, "r") as f:
                metadata["total_rows"] = sum(1 for _ in f) - 1

        except ImportError:
            # Fallback without pandas
            import csv

            with open(file_path, "r", newline="") as f:
                reader = csv.reader(f)
                headers = next(reader, [])
                row_count = sum(1 for _ in reader)
            metadata["columns"] = headers
            metadata["column_count"] = len(headers)
            metadata["total_rows"] = row_count

        except Exception as e:
            metadata["error"] = str(e)

        return metadata

    # ── XLSX ─────────────────────────────────

    def _extract_xlsx_metadata(
        self, file_path: str
    ) -> dict[str, Any]:
        metadata: dict[str, Any] = {"meta_type": "document"}

        try:
            import openpyxl

            wb = openpyxl.load_workbook(
                file_path, read_only=True, data_only=True
            )
            metadata["sheet_names"] = wb.sheetnames
            metadata["sheet_count"] = len(wb.sheetnames)

            props = wb.properties
            if props:
                metadata["author"] = props.creator
                metadata["title"] = props.title
                metadata["creation_date"] = (
                    props.created.isoformat() if props.created else None
                )
                metadata["modification_date"] = (
                    props.modified.isoformat() if props.modified else None
                )

            # First sheet dimensions
            ws = wb.active
            if ws:
                metadata["active_sheet_rows"] = ws.max_row
                metadata["active_sheet_cols"] = ws.max_column

            wb.close()

        except ImportError:
            metadata["warning"] = "openpyxl not installed"
        except Exception as e:
            metadata["error"] = str(e)

        return metadata
